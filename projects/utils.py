import re
from datetime import date
from pprint import pprint

from django.db.models import Q, QuerySet
from django.http import FileResponse
from docx import Document
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.table import Table

from projects.models import Projects, KindOfOfInfluences, ViolatorLvls, Bdus, ObjectOfInfluences, Capecs, \
    NormativeDocuments, Components, SPMethods, Tactics


def add_bullet_list(paragraph, is_text, list_items):
    for item in list_items:
        run = paragraph.add_run(f'• {item}')
        if is_text:
            run.add_break(WD_BREAK.LINE)


def generate_doc(project: Projects):
    neg_pos = genereate_neg_con_table(project)
    doc: Document = Document('shablon_modeli_ugroz.docx')
    for paragraph in doc.paragraphs:
        paragraph.text = paragraph.text.replace('__должность__', project.r_persons.all()[0].appointment)
        paragraph.text = paragraph.text.replace('__название_организации__',
                                                'НАЗВАНИЕ ОРГАНИЗАЦИИ, ФУНКЦИЯ ТРЕБУЕТ ДОРАБОТКИ')
        paragraph.text = paragraph.text.replace('__ФИО__', project.r_persons.all()[0].name)
        paragraph.text = paragraph.text.replace('__дата__', date.today().strftime('%d.%m.%Y'))
        paragraph.text = paragraph.text.replace('__тип_системы__', project.get_type_display())
        paragraph.text = paragraph.text.replace('__название_системы__', project.name_project)
        paragraph.text = paragraph.text.replace('__описание__системы__',
                                                project.description if project.description is not None else "описание отсутствует")
        paragraph.text = paragraph.text.replace('__год__', str(date.today().year))
        if "__название_нег_поз__" in paragraph.text:
            paragraph.text = paragraph.text.replace('__название_нег_поз__', '')
            add_bullet_list(paragraph, True,
                            project.negative_consequences.all().order_by('id').values_list('name', flat=True))
        if '__список_нормативки_для_соот_системы__' in paragraph.text:
            paragraph.text = paragraph.text.replace('__список_нормативки_для_соот_системы__', '')
            add_bullet_list(paragraph, True,
                            NormativeDocuments.objects.all().
                            filter(type=project.type).
                            order_by('id').
                            values_list('name', flat=True))
        if '__список_СП__' in paragraph.text:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            paragraph.text = paragraph.text.replace('__список_СП__', '')
            components = project.components.all()
            smp: QuerySet = SPMethods.objects.none()
            for component in components:
                smp = smp | component.spmethods.all()
            smps = smp.distinct().values_list('name', 'alias')
            add_bullet_list(paragraph, True, [f'{smp[1]}: {smp[0]}' for smp in smps])
    # автозаполнения даты
    table = doc.tables[0]
    current_date = str(date.today())
    table.cell(0, 1).text = current_date
    table2 = doc.tables[1]
    # заполнение таблицы нег.последствия
    for key in neg_pos:
        for elems in neg_pos[key]:  # тут мы перебираем список
            new_row = table2.add_row().cells
            new_row[0].text = key
            new_row[1].text = elems
    # удаляю повтор.значения в таблице
    row_count = len(table2.rows)
    col_count = len(table2.columns)
    listok = []
    for row in range(row_count):
        for col in range(col_count):
            if table2.cell(row, col).text in listok:
                table2.cell(row, col).text = ""
            else:
                listok.append((table2.cell(row, col).text))
    # скрещиваю пустые ячейки с пред идущимиЫ
    for row in range(row_count):
        if table2.cell(row, 0).text == '':
            table2.cell(row - 1, 0).merge(table2.cell(row, 0))
    # работа с таблицей №3   {         {     {       }        }          }
    objinftable = generate_obj_inf_table(project)
    table3 = doc.tables[2]
    for key in objinftable:
        for elem in objinftable[key]:
            # for key_to in objinftable[key][elem]:
            new_row = table3.add_row().cells
            new_row[0].text = key
            new_row[1].text = elem
            new_row[2].text = objinftable[key][elem]
    # работа с таблицей №4
    gen_vio = generate_violators_type_table(project)
    table4 = doc.tables[3]
    for key in gen_vio:
        for elem in gen_vio[key]:
            new_row = table4.add_row().cells
            new_row[0].text = str(key)
            new_row[1].text = str(elem)
            new_row[2].text = str(gen_vio[key][elem])
    # работа с таблицей №5(потенциал нарушителя)
    gen_poten = generate_violators_potential_table(project)
    table5: Table = doc.tables[4]
    for key in gen_poten:
        for elem in gen_poten[key]:
            new_row = table5.add_row().cells
            new_row[0].text = key
            new_row[1].text = elem
            new_row[2].text = str(gen_poten[key][elem])
    row_count = len(table5.rows)
    col_count = len(table5.columns)
    listok = []
    for row in range(row_count):
        for col in range(col_count):
            if table5.cell(row, col).text in listok and table5.cell(row, col).text not in ['высокий', 'низкий',
                                                                                           'средний']:
                table5.cell(row, col).text = ""
            else:
                listok.append((table5.cell(row, col).text))
    for row in range(row_count):
        if table5.cell(row, 0).text == '':
            table5.cell(row - 1, 0).merge(table5.cell(row, 0))
    # работа с таблицей №6(способы реализации)
    print('66666666666666666666666666666666666')
    gen_spm = generate_sp_methods_table(project)
    table6: Table = doc.tables[5]
    for obj in gen_spm:
        for component in gen_spm[obj]:
            for spm in gen_spm[obj][component]:
                new_row = table6.add_row().cells
                new_row[0].text = obj
                new_row[1].text = component
                new_row[2].text = spm
    table6 = clear_duplicate(table6)

    # работа с таблицей №7(потенциал нарушителя)

    print('777777777777777777777777777777777777777777777')
    gen_pot_bdu = generate_potential_bdus_table(project)
    smps_set = set()
    table7: Table = doc.tables[6]
    for number_bdus in gen_pot_bdu:
        for name_bdus in gen_pot_bdu[number_bdus]:
            for vector_capec in gen_pot_bdu[number_bdus][name_bdus]:
                for neg_con in gen_pot_bdu[number_bdus][name_bdus][vector_capec]:
                    for violator in gen_pot_bdu[number_bdus][name_bdus][vector_capec][neg_con]:
                        for obj_of_imp in gen_pot_bdu[number_bdus][name_bdus][vector_capec][neg_con][violator]:
                            for component_t in gen_pot_bdu[number_bdus][name_bdus][vector_capec][neg_con][violator][
                                obj_of_imp]:
                                for spm_t in \
                                        gen_pot_bdu[number_bdus][name_bdus][vector_capec][neg_con][violator][
                                            obj_of_imp][
                                            component_t]:
                                    new_row = table7.add_row().cells
                                    new_row[0].text = number_bdus
                                    new_row[1].text = name_bdus
                                    new_row[2].text = vector_capec
                                    new_row[3].text = neg_con
                                    new_row[4].text = violator
                                    new_row[5].text = obj_of_imp
                                    new_row[6].text = component_t
                                    new_row[7].text = spm_t
                                    smps_set.add(spm_t)
    table7 = clear_duplicate(table7)

    # работа с таблицей №8 (cпособы реализации и их техники)
    print('8888888888888888888888888888888888888888888888')
    table8: Table = doc.tables[7]
    gen_tactics = generate_tactic_table(smps_set)
    # работа с таблицей №9  (актуальные угрозы)
    # todo сместить номер таблицы на +2

    # gen_bdu = generate_bdu_table(project)
    # print('88888888888888888888888888888888888888888888')
    # table7: Table = doc.tables[6]
    # for number_ugroz in gen_bdu:
    #     print(number_ugroz)
    #     for name_ugroz in gen_bdu[number_ugroz]:
    #         for uyazvimost in gen_bdu[number_ugroz][name_ugroz]:
    #             for vectorCapec in gen_bdu[number_ugroz][name_ugroz][uyazvimost]:
    #                 for negativ in gen_bdu[number_ugroz][name_ugroz][uyazvimost][vectorCapec]:
    #                     for object in gen_bdu[number_ugroz][name_ugroz][uyazvimost][vectorCapec][negativ]:
    #                         for tn in gen_bdu[number_ugroz][name_ugroz][uyazvimost][vectorCapec][negativ][object]:
    #                             new_row = table7.add_row().cells
    #                             new_row[0].text = number_ugroz
    #                             new_row[1].text = name_ugroz
    #                             new_row[2].text = uyazvimost
    #                             new_row[3].text = vectorCapec
    #                             new_row[4].text = negativ
    #                             new_row[5].text = object
    #                             new_row[6].text = tn
    #                             new_row[7].text = \
    #                                 gen_bdu[number_ugroz][name_ugroz][uyazvimost][vectorCapec][negativ][object][tn]
    # table7 = clear_duplicate(table7)
    doc.save('новое_имя_файла.docx')
    word_file_path = 'новое_имя_файла.docx'
    response = FileResponse(open(word_file_path, 'rb'))
    response['Content-Disposition'] = 'attachment; filename="новое_имя_файла.docx"'
    return response


def clear_duplicate(table):
    row_count = len(table.rows)
    col_count = len(table.columns)
    listok = []
    temp = None
    for row in range(row_count):
        if table.cell(row, 0).text != temp:
            listok.clear()
            temp = table.cell(row, 0).text
            print(temp)
        for col in range(col_count):
            if table.cell(row, col).text in listok:
                table.cell(row, col).text = ""
            else:
                listok.append((table.cell(row, col).text))
    for row in range(row_count):
        for col in range(col_count):
            if table.cell(row, col).text == '':
                table.cell(row - 1, col).merge(table.cell(row, col))
    return table


def generate_tactic_table(smps: set[str]):
    smps = sorted(smps)
    result = {}
    for smp in smps:
        smp = SPMethods.objects.get(alias=smp)
        tactics = smp.tactics.all()
        if tactics.exists():
            for tactic in tactics:
                if smp.alias in result:
                    temp = result[smp.alias]
                    result.update({smp.alias: temp | {tactic.alias: [tactic.name]}})
                else:
                    result[smp.alias] = {tactic.alias: [tactic.name]}
    return result


def generate_sp_methods_table(project):
    table = {}
    actual_lvl = project.violators.all().values_list('lvl_id', flat=True).distinct()
    actual_lvl = ['В.' + str(i) for i in actual_lvl]
    objs = project.object_inf.all().order_by('id')
    pr_components = project.components.all()
    for obj in objs:
        components: QuerySet[Components] = obj.components.all() & pr_components
        for component in components:
            spmethods = component.spmethods.all() if component.spmethods.all().exists() else [
                'СП отсутствует']
            for spm in spmethods:
                if type(spm) is str:
                    spm_name = spm
                elif spm.violator_lvls in actual_lvl:
                    spm_name = spm.name
                else:
                    spm_name = 'СП отсутствует, так как нет подходящих нарушителей'
                if f'{obj.name} ({obj.alias})' in table:
                    temp = table[f'{obj.name} ({obj.alias})']
                    table.update({f'{obj.name} ({obj.alias})': temp | {component.name: [spm_name]}})
                else:
                    table[f'{obj.name} ({obj.alias})'] = {component.name: [spm_name]}

    # todo создать excel файл и вернуть его
    return table


def genereate_neg_con_table(project: Projects):
    table = {}
    neg_cons = project.negative_consequences.all().order_by('type')
    for neg_con in neg_cons:
        if neg_con.type in table:
            table[neg_con.type].append(neg_con.name)
        else:
            table[neg_con.type] = [neg_con.name]
    # pprint(table)

    # todo создать excel файл и вернуть его
    return table


def generate_obj_inf_table(project: Projects):
    table = {}
    neg_cons = project.negative_consequences.all()
    for neg_con in neg_cons:
        if 'физическому' in neg_con.type:
            lvl = 'У1'
        elif 'юридическому' in neg_con.type:
            lvl = 'У2'
        else:
            lvl = 'У3'
        objs = neg_con.object_of_influence.all()
        for obj in objs:
            if obj in project.object_inf.all():
                kind = KindOfOfInfluences.objects.get(object_of_inf=obj, neg_cons=neg_con).kind_of_inf
                if f'{neg_con.name} ({lvl})' in table:
                    temp = table[f'{neg_con.name} ({lvl})']
                    table.update({f'{neg_con.name} ({lvl})': temp | {obj.name: [kind]}})
                else:
                    table[f'{neg_con.name} ({lvl})'] = {obj.name: [kind]}
        if f'{neg_con.name} ({lvl})' not in table:
            table[f'{neg_con.name} ({lvl})'] = {
                'нет потенциальных объектов воздействия': 'воздействие отсутствует'}

            # todo создать excel файл и вернуть его
    # print(table)
    return table


def generate_violators_type_table(project: Projects):
    # вид нарушителя(название), тип нарушителя (внеш, внут), Возможные цели (мотивация) реализации угроз безопасности информации
    table = {}
    violators = project.violators.all()
    for violator in violators:
        if violator.type == 1:
            types = ['внутренний']
        elif violator.type == 2:
            types = ['внешний']
        else:
            types = ['внутренний', 'внешний']
        for type in types:
            if violator.name in table:
                temp = table[violator.name]
                table.update({violator.name: temp | {type: violator.motives}})
            else:
                table[violator.name] = {type: violator.motives}
    # todo создать excel файл и вернуть его
    # pprint(table)
    return table


def generate_violators_potential_table(project: Projects):
    # Уровень возможностей, вид нарушителя(название), потенциал нарушителя
    table = {}
    vlvls = ViolatorLvls.objects.all().order_by('lvl')
    for lvl in vlvls:
        violators = lvl.violators.all()
        for violator in violators:
            if violator.potential == 1:
                potential = 'низкий'
            elif violator.type == 2:
                potential = 'средний'
            else:
                potential = 'высокий'
            if violator in project.violators.all():
                if lvl.name in table:
                    temp = table[lvl.name]
                    table.update({lvl.name: temp | {violator.name: potential}})
                else:
                    table[lvl.name] = {violator.name: potential}
    # todo создать excel файл и вернуть его
    # pprint(table)
    return table


def generate_bdu_table(project: Projects):
    '''
    эта таблица не уместится в обычном варианте ворда,
    лист нужно будет развернуть горизонтально что бы она поместилась хоть как-то
    и придётся уменьшить шрифт для неё
    :param project:
    :return:
    '''
    # номер угрозы, название, уязвимость(опционально), вектор капек, нег пос, объект воздействия, нарушитель, сценарий реализации
    table = {}
    correct_obj_list = project.object_inf.all()

    bdus: QuerySet[Bdus] = form_bdus_list_for(project).order_by('id').all()
    # Bdus.objects.all()
    bdus.all()
    for bdu in bdus:
        capecs: QuerySet[Capecs] = bdu.capecs.all()
        for capec in capecs:
            number = f"УБИ.{bdu.id}"

            # формирование цепочки капек от родителя до текущего
            cpc = form_capec_vector_for(capec)

            # нег взять дискрипшион  обрезать строку "Угроза заключается в возможности"
            # включительно убрать _x000D и всё что после
            neg_con = bdu.description.replace("Угроза заключается в возможности", '')
            neg_con = neg_con.replace(r'_x000D*', '')
            neg_con = re.sub('(_x000D).*', '', neg_con, re.DOTALL).split('\n', 1)[0]

            vulnerability = 'ручное заполнение'
            obj_of_infs = bdu.bdus.all()
            violator = bdu.violator
            scenario = 'сценарий реализации - функция в разработке'
            violator_dict = {violator: scenario}
            neg_con_dict = {}
            obj_of_infs = obj_of_infs & correct_obj_list
            if obj_of_infs:
                for obj_of_inf in obj_of_infs:
                    obj_of_inf = obj_of_inf.name
                    if neg_con in neg_con_dict:
                        temp = neg_con_dict[neg_con]
                        neg_con_dict.update({neg_con: temp | {obj_of_inf: violator_dict}})
                    else:
                        neg_con_dict[neg_con] = {obj_of_inf: violator_dict}
            else:
                obj_of_inf = 'нет актуальных объектов'
                if neg_con in neg_con_dict:
                    temp = neg_con_dict[neg_con]
                    neg_con_dict.update({neg_con: temp | {obj_of_inf: violator_dict}})
                else:
                    neg_con_dict[neg_con] = {obj_of_inf: violator_dict}
            vulnerability_dict = {vulnerability: neg_con_dict}
            cpc_dict = {cpc: vulnerability_dict}
            if number in table:
                temp = table[number]
                table.update({number: temp | {bdu.name: cpc_dict}})
            else:
                table[number] = {bdu.name: cpc_dict}

    # todo создать excel файл и вернуть его

    pprint(table)
    return table


def generate_potential_bdus_table(project: Projects):
    '''
       эта таблица не уместится в обычном варианте ворда,
       лист нужно будет развернуть горизонтально что бы она поместилась хоть как-то
       и придётся уменьшить шрифт для неё
       :param project:
       :return:
       '''
    # номер угрозы, название, вектор капек, нег пос,нарушитель, объект воздействия, компонент, способ реализации
    table = {}
    correct_obj_list = project.object_inf.all()

    bdus: QuerySet[Bdus] = form_bdus_list_for(project).order_by('id').all()
    # Bdus.objects.all()
    bdus.all()
    for bdu in bdus:
        capecs: QuerySet[Capecs] = bdu.capecs.all()
        for capec in capecs:
            number = f"УБИ.{bdu.id}"

            # формирование цепочки капек от родителя до текущего
            cpc = form_capec_vector_for(capec)

            # нег взять дискрипшион  обрезать строку "Угроза заключается в возможности"
            # включительно убрать _x000D и всё что после
            neg_con = bdu.description.replace("Угроза заключается в возможности", '')
            neg_con = neg_con.replace(r'_x000D*', '')
            neg_con = re.sub('(_x000D).*', '', neg_con, re.DOTALL).split('\n', 1)[0]
            violator = bdu.violator
            violator_dict = {}
            obj_of_infs: QuerySet[ObjectOfInfluences] = bdu.bdus.all() & correct_obj_list
            if obj_of_infs.exists():
                obj_of_infs_dict = {}
                for obj_of_inf in obj_of_infs:
                    components = obj_of_inf.components.all() & project.components.all()
                    if components.exists():
                        components_dict = {}
                        for component in components:
                            spms = component.spmethods.all()
                            # подвязка сп к компоненту
                            if spms.exists():
                                for spm in spms:
                                    if component.alias in components_dict:
                                        # возможный проёб
                                        components_dict[component.alias].append(spm.alias)
                                    else:
                                        components_dict[component.alias] = [spm.alias]
                            else:
                                components_dict[component.alias] = ['нет способов реализации']

                            # подвязка компонента к объекту
                            if obj_of_inf.alias in obj_of_infs_dict:
                                temp = obj_of_infs_dict[obj_of_inf.alias]
                                obj_of_infs_dict.update({obj_of_inf.alias: temp | components_dict})
                            else:
                                obj_of_infs_dict[obj_of_inf.alias] = components_dict
                    else:
                        obj_of_infs_dict[obj_of_inf.alias] = {'компоненты отсутствуют': ['нет способов реализации']}
                violator_dict[violator] = obj_of_infs_dict
            else:
                violator_dict[violator] = {
                    'объекты отсутствуют': {'компоненты отсутствуют': ['нет способов реализации']}}
            neg_con_dict = {neg_con: violator_dict}
            cpc_dict = {cpc: neg_con_dict}
            if number in table:
                temp = table[number]
                table.update({number: temp | {bdu.name: cpc_dict}})
            else:
                table[number] = {bdu.name: cpc_dict}
    # todo создать excel файл и вернуть его

    # pprint(table)
    return table


def form_capec_vector_for(capec: Capecs) -> str:
    cpc: list[str] = []
    while True:
        cpc.append(f'CAPEC-{capec.id}: {capec.name}')
        # todo должно быть is None, но база данных хранит не все капеки, поэтому алгоритм кидает ошибки
        if (capec.parent_id is not None):
            break
        capec = Capecs.objects.get(id=capec.parent_id)

    result = '\n'.join(cpc)
    return result


def form_bdus_list_for(project: Projects) -> QuerySet[Bdus]:
    # функция которая подвязывает к проекту актуальные бдухи
    # фильтрация по объектам
    objects = ObjectOfInfluences.objects.filter(name__in=project.object_inf.values_list('name', flat=True))
    bdu = Bdus.objects.none()
    for object in objects:
        a = object.bdus.all()
        bdu = bdu | a
    # при фильтрации потегам селект none or true/false
    bdu = bdu.distinct()
    bdu = bdu.filter(Q(is_wireless=project.is_wireless) | Q(is_wireless=None))
    bdu = bdu.filter(Q(is_grid=project.is_grid) | Q(is_grid=None))
    bdu = bdu.filter(Q(is_virtual=project.is_virtual) | Q(is_virtual=None))
    bdu = bdu.filter(Q(is_cloud=project.is_cloud) | Q(is_cloud=None)).order_by('id')

    # фильтрация по нарушителям
    # todo я переделал функцию, она сломалась, нужно затестить
    # violators = project.get_violator_lvl_names()
    # bdu = bdu.filter(violator__in=violators)

    return bdu
